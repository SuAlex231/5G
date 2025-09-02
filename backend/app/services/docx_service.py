from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import io
from typing import List
import uuid

from ..models.models import Ticket, Image
from ..core.config import settings
from .minio_service import MinIOService


async def create_ticket_docx(ticket: Ticket, images: List[Image]) -> str:
    """Create a DOCX document with ticket images."""
    doc = Document()
    minio_service = MinIOService()
    
    # Add title
    title = doc.add_heading(f'Ticket: {ticket.title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add ticket info
    doc.add_heading('Ticket Information', level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Ticket Number', ticket.ticket_number),
        ('Status', ticket.status.title()),
        ('Priority', ticket.priority.title()),
        ('Created', ticket.created_at.strftime('%Y-%m-%d %H:%M'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)
    
    # Add form data if available
    if ticket.form_data:
        doc.add_heading('Form Data', level=1)
        for key, value in ticket.form_data.items():
            doc.add_paragraph(f'{key}: {value}')
    
    # Add images section
    if images:
        doc.add_heading('Images', level=1)
        
        for i, image in enumerate(images):
            try:
                # Get image content from MinIO
                image_content = await minio_service.get_file_content(image.file_path)
                
                # Open image with PIL to get dimensions
                pil_image = PILImage.open(io.BytesIO(image_content))
                width, height = pil_image.size
                
                # Calculate scaling to fit page width (6 inches max)
                max_width = 6.0
                if width > height:
                    # Landscape orientation
                    scale = min(max_width / (width / 100), 4.0)
                else:
                    # Portrait orientation
                    scale = min(max_width / (width / 100), 6.0)
                
                # Add image to document
                doc.add_paragraph(f'Image {i + 1}: {image.original_filename}')
                
                # Save image temporarily for docx
                temp_image_stream = io.BytesIO(image_content)
                doc.add_picture(temp_image_stream, width=Inches(scale))
                
                # Add page break after each image except the last one
                if i < len(images) - 1:
                    doc.add_page_break()
                    
            except Exception as e:
                # If image can't be processed, add error note
                doc.add_paragraph(f'Image {i + 1}: {image.original_filename} (Error loading: {str(e)})')
    
    # Save document to MinIO
    doc_filename = f"exports/ticket_{ticket.id}_images_{uuid.uuid4().hex[:8]}.docx"
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    # Upload to MinIO
    try:
        minio_service.client.put_object(
            settings.MINIO_BUCKET_EXPORTS,
            doc_filename,
            doc_stream,
            length=doc_stream.getbuffer().nbytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Generate download URL
        download_url = minio_service.get_presigned_download_url(
            doc_filename,
            settings.MINIO_BUCKET_EXPORTS,
            expires_in=3600  # 1 hour
        )
        
        return download_url
        
    except Exception as e:
        raise Exception(f"Failed to save DOCX file: {e}")


async def create_bulk_docx(tickets: List[Ticket]) -> str:
    """Create a DOCX document with multiple tickets (for bulk export)."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Ticket Export Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    doc.add_heading('Export Summary', level=1)
    doc.add_paragraph(f'Total tickets: {len(tickets)}')
    doc.add_paragraph(f'Export date: {ticket.created_at.strftime("%Y-%m-%d %H:%M")}')
    
    # Add each ticket
    for i, ticket in enumerate(tickets):
        if i > 0:
            doc.add_page_break()
            
        # Ticket header
        doc.add_heading(f'Ticket #{ticket.ticket_number}', level=1)
        
        # Basic info
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Title', ticket.title),
            ('Status', ticket.status.title()),
            ('Priority', ticket.priority.title()),
            ('Created', ticket.created_at.strftime('%Y-%m-%d %H:%M'))
        ]
        
        for j, (label, value) in enumerate(info_data):
            info_table.cell(j, 0).text = label
            info_table.cell(j, 1).text = str(value)
        
        # Form data
        if ticket.form_data:
            doc.add_heading('Details', level=2)
            for key, value in ticket.form_data.items():
                doc.add_paragraph(f'{key}: {value}')
    
    # Save and return download URL (similar to single ticket)
    doc_filename = f"exports/bulk_tickets_{uuid.uuid4().hex[:8]}.docx"
    # ... rest of save logic
    
    return doc_filename